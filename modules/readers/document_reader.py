"""
╔══════════════════════════════════════════════════════════════════════════════╗
║           modules/readers/document_reader.py  (V1.3)                       ║
║           Responsabilidade: Leitura de arquivos (TXT / DOCX / PDF)         ║
╚══════════════════════════════════════════════════════════════════════════════╝

MUDANÇA V1.2 → V1.3:
    - Retorno padronizado com ok() / erro()
    - Tratamento robusto de encoding para TXT
    - Extração de tabelas DOCX melhorada
    - PDF: limite de páginas configurável (evita PDFs gigantes)
    - Logging mais claro
"""

import os
import logging
import fitz                      # PyMuPDF — leitura de PDF
from docx import Document        # python-docx — leitura de Word
from modules.utils.response import ok, erro

logger = logging.getLogger(__name__)

# Extensões suportadas
EXTENSOES = {".txt", ".docx", ".pdf"}

# Limite de páginas por PDF (evita documentos de 300 páginas travarem o sistema)
# Aumente se precisar processar PDFs muito longos
MAX_PAGINAS_PDF = 40


class DocumentReader:
    """
    Lê qualquer arquivo suportado e retorna texto puro.

    Uso:
        reader = DocumentReader()
        resultado = reader.ler("Cursos/Dev/Backend/UC_Backend.pdf")

        if resultado["sucesso"]:
            texto = resultado["dados"]["texto"]
    """

    def ler(self, caminho: str) -> dict:
        """
        Lê um arquivo e retorna seu conteúdo como texto.

        Returns:
            ok(texto="...", paginas=N, extensao=".pdf")
            erro("mensagem de erro")
        """
        if not caminho:
            return erro("Caminho não fornecido")

        if not os.path.exists(caminho):
            return erro(f"Arquivo não encontrado: {caminho}")

        _, ext = os.path.splitext(caminho)
        ext = ext.lower()

        if ext not in EXTENSOES:
            return erro(f"Formato não suportado: '{ext}'. Use: {', '.join(EXTENSOES)}")

        try:
            if ext == ".txt":
                return self._ler_txt(caminho)
            elif ext == ".docx":
                return self._ler_docx(caminho)
            elif ext == ".pdf":
                return self._ler_pdf(caminho)
        except Exception as e:
            logger.error(f"Erro ao ler '{caminho}': {e}")
            return erro(f"Erro ao processar arquivo: {str(e)}")

    def ler_varios(self, caminhos: list) -> dict:
        """
        Lê múltiplos arquivos e consolida o texto.
        Ignora arquivos que falham — não trava o sistema.

        Returns:
            ok(texto="...", lidos=N, falhas=N, detalhes=[...])
        """
        textos      = []
        lidos       = 0
        falhas      = 0
        detalhes    = []

        for caminho in caminhos:
            resultado = self.ler(caminho)
            nome = os.path.basename(caminho)

            if resultado["sucesso"]:
                texto = resultado["dados"].get("texto", "")
                if texto.strip():
                    textos.append(f"\n── {nome} ──\n{texto}")
                    lidos += 1
                    detalhes.append({"arquivo": nome, "status": "ok"})
            else:
                falhas += 1
                detalhes.append({
                    "arquivo": nome,
                    "status": "erro",
                    "mensagem": resultado.get("erro", "")
                })
                logger.warning(f"Ignorando '{nome}': {resultado.get('erro')}")

        return ok(
            texto     = "\n\n".join(textos),
            lidos     = lidos,
            falhas    = falhas,
            detalhes  = detalhes,
        )

    # ─── Leitores por formato ─────────────────────────────────────────────────

    def _ler_txt(self, caminho: str) -> dict:
        """Lê TXT tentando UTF-8 e fazendo fallback para latin-1."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(caminho, "r", encoding=encoding) as f:
                    texto = f.read()
                logger.info(f"TXT lido ({encoding}): {caminho} — {len(texto)} chars")
                return ok(texto=texto, extensao=".txt", paginas=1)
            except UnicodeDecodeError:
                continue
        return erro(f"Não foi possível decodificar o arquivo TXT: {caminho}")

    def _ler_docx(self, caminho: str) -> dict:
        """
        Lê DOCX extraindo parágrafos e conteúdo de tabelas.
        Tabelas são comuns em documentos SENAI (competências, capacidades).
        """
        doc    = Document(caminho)
        partes = []

        # Parágrafos normais
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                partes.append(t)

        # Tabelas (competências, capacidades, critérios avaliativos)
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                if celulas:
                    partes.append(" | ".join(celulas))

        texto = "\n".join(partes)
        logger.info(f"DOCX lido: {caminho} — {len(texto)} chars")
        return ok(texto=texto, extensao=".docx", paginas=len(doc.paragraphs))

    def _ler_pdf(self, caminho: str) -> dict:
        """
        Lê PDF página a página com limite de segurança.

        MAX_PAGINAS_PDF evita que PDFs gigantes (provas antigas, apostilas)
        consumam tokens desnecessários. O extrator vai resumir o conteúdo
        antes de enviar para a IA de qualquer forma.
        """
        partes        = []
        total_paginas = 0

        with fitz.open(caminho) as pdf:
            total_paginas = len(pdf)
            paginas_ler   = min(total_paginas, MAX_PAGINAS_PDF)

            if total_paginas > MAX_PAGINAS_PDF:
                logger.warning(
                    f"PDF com {total_paginas} páginas — lendo apenas as "
                    f"primeiras {MAX_PAGINAS_PDF}: {caminho}"
                )

            for i in range(paginas_ler):
                texto_pag = pdf[i].get_text("text")
                if texto_pag.strip():
                    partes.append(texto_pag)

        texto = "\n".join(partes)
        logger.info(
            f"PDF lido: {caminho} — "
            f"{paginas_ler}/{total_paginas} págs — {len(texto)} chars"
        )
        return ok(
            texto         = texto,
            extensao      = ".pdf",
            paginas       = paginas_ler,
            total_paginas = total_paginas,
            truncado      = total_paginas > MAX_PAGINAS_PDF,
        )
