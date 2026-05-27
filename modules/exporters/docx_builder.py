"""
╔══════════════════════════════════════════════════════════════════════════════╗
║         modules/exporters/docx_builder.py  (V1.8 — NOVO)                   ║
║         Constrói DOCX profissional a partir do JSON pedagógico             ║
╚══════════════════════════════════════════════════════════════════════════════╝

Responsabilidade:
    Receber os capítulos em formato JSON (gerados pelo Gemini)
    e construir um arquivo .docx profissional com:
        - Capa com identidade SENAI
        - Sumário automático
        - Capítulos com hierarquia visual
        - Seções com estilos consistentes
        - Blocos de código formatados
        - Tabelas estilizadas
        - Boxes de destaque (Atenção, Dica, SAEP, Prática)
        - Rodapé com numeração de páginas
        - Metadados do documento

Por que python-docx em vez de OpenAI para formatar:
    - Determinístico: o resultado é sempre consistente
    - Custo zero: sem API paga para formatação
    - Mais rápido: sem latência de rede adicional
    - Mais controle: definimos exatamente cada estilo
"""

import os
import re
import logging
from datetime import datetime
from typing import Optional, List

from docx import Document
from docx.shared import (
    Pt, Cm, RGBColor, Inches
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modules.utils.response import ok, erro

logger = logging.getLogger(__name__)

# ─── Paleta de cores SENAI ────────────────────────────────────────────────────
AZUL_ESCURO  = RGBColor(0x0d, 0x3b, 0x66)
AZUL_MEDIO   = RGBColor(0x1a, 0x5f, 0xa8)
AZUL_CLARO   = RGBColor(0xea, 0xf2, 0xfb)
TEXTO_PRIN   = RGBColor(0x1a, 0x2a, 0x3a)
TEXTO_SEC    = RGBColor(0x4a, 0x6a, 0x8a)
CINZA_CODIGO = RGBColor(0x1e, 0x29, 0x3b)
CINZA_CLARO  = RGBColor(0xf5, 0xf8, 0xff)
VERDE        = RGBColor(0x0f, 0x6b, 0x3a)
VERDE_FUNDO  = RGBColor(0xe8, 0xf5, 0xee)
AMARELO      = RGBColor(0xd9, 0x77, 0x06)
AMARELO_FUND = RGBColor(0xff, 0xfb, 0xeb)
ROXO         = RGBColor(0x7c, 0x3a, 0xed)
ROXO_FUNDO   = RGBColor(0xf3, 0xf0, 0xff)
LARANJA      = RGBColor(0xea, 0x58, 0x0c)
LARANJA_FUND = RGBColor(0xff, 0xf7, 0xed)
BRANCO       = RGBColor(0xff, 0xff, 0xff)


class DocxBuilder:
    """
    Constrói um arquivo DOCX profissional a partir dos dados JSON dos capítulos.

    Uso:
        builder = DocxBuilder(pasta_output="output")
        resultado = builder.construir(
            titulo_apostila="...",
            curso="...", uc="...", bloco_aula="...",
            capitulos=[{JSON do cap 1}, {JSON do cap 2}, ...],
            metricas={...}
        )
    """

    def __init__(self, pasta_output: str = "output"):
        self.pasta_output = pasta_output
        os.makedirs(pasta_output, exist_ok=True)

    def construir(
        self,
        titulo_apostila:  str,
        curso:            str,
        uc:               str,
        bloco_aula:       str,
        capitulos:        list,
        metricas:         dict = None,
        professor:        Optional[str] = None,
    ) -> dict:
        """
        Constrói o DOCX completo.

        Args:
            titulo_apostila: Título gerado pelo mapa pedagógico
            curso/uc/bloco:  Dados educacionais
            capitulos:       Lista de dicts com JSON de cada capítulo
            metricas:        Métricas de geração (tokens, palavras, tempo)
            professor:       Nome do professor (para metadados)

        Returns:
            ok(nome_arquivo, caminho, tamanho, data_geracao)
            erro("mensagem")
        """
        try:
            doc = Document()
            self._configurar_documento(doc)
            self._definir_estilos(doc)

            # ── Capa ──────────────────────────────────────────────────────────
            self._adicionar_capa(doc, titulo_apostila, curso, uc, bloco_aula,
                                 len(capitulos), professor)

            # ── Sumário ───────────────────────────────────────────────────────
            self._adicionar_sumario(doc, capitulos)

            # ── Capítulos ─────────────────────────────────────────────────────
            for cap_dados in capitulos:
                self._adicionar_capitulo(doc, cap_dados)

            # ── Rodapé com numeração ──────────────────────────────────────────
            self._adicionar_rodape(doc, uc, bloco_aula)

            # ── Salva ─────────────────────────────────────────────────────────
            nome    = self._gerar_nome(curso, uc, bloco_aula)
            caminho = os.path.join(self.pasta_output, nome)
            doc.save(caminho)

            tamanho = os.path.getsize(caminho)
            data_hr = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

            logger.info(f"DOCX gerado: {nome} ({self._fmt(tamanho)})")

            return ok(
                nome_arquivo    = nome,
                caminho_arquivo = caminho,
                tamanho_bytes   = tamanho,
                tamanho_legivel = self._fmt(tamanho),
                data_geracao    = data_hr,
            )

        except Exception as e:
            logger.error(f"Erro ao construir DOCX: {e}", exc_info=True)
            return erro(f"Falha ao construir DOCX: {str(e)}")

    def listar(self) -> list:
        """Lista apostilas DOCX geradas."""
        if not os.path.exists(self.pasta_output):
            return []
        apostilas = []
        for nome in os.listdir(self.pasta_output):
            if not nome.endswith(".docx"):
                continue
            caminho = os.path.join(self.pasta_output, nome)
            stat    = os.stat(caminho)
            apostilas.append({
                "nome_arquivo": nome,
                "caminho":      caminho,
                "tamanho":      self._fmt(stat.st_size),
                "data_criacao": datetime.fromtimestamp(stat.st_ctime).strftime("%d/%m/%Y %H:%M"),
                "tipo":         "docx",
            })
        return sorted(apostilas, key=lambda x: x["data_criacao"], reverse=True)

    def deletar(self, nome_arquivo: str) -> dict:
        """Remove um DOCX do disco."""
        nome_seguro = os.path.basename(nome_arquivo)
        if not nome_seguro.endswith(".docx"):
            return erro("Apenas arquivos .docx podem ser removidos")
        caminho = os.path.join(self.pasta_output, nome_seguro)
        if not os.path.exists(caminho):
            return erro(f"Arquivo não encontrado: {nome_seguro}")
        try:
            os.remove(caminho)
            logger.info(f"DOCX removido: {nome_seguro}")
            return ok(mensagem=f"'{nome_seguro}' removida com sucesso")
        except Exception as e:
            return erro(f"Falha ao remover: {e}")

    # ═══════════════════════════════════════════════════════════
    # CONFIGURAÇÃO DO DOCUMENTO
    # ═══════════════════════════════════════════════════════════

    def _configurar_documento(self, doc: Document):
        """Configura margens e tamanho de página."""
        from docx.shared import Cm
        sec = doc.sections[0]
        sec.page_width    = Inches(8.27)   # A4
        sec.page_height   = Inches(11.69)  # A4
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

        # Metadados do documento
        doc.core_properties.author   = "Gerador de Apostilas SENAI V1.8"
        doc.core_properties.company  = "SENAI"
        doc.core_properties.created  = datetime.now()

    def _definir_estilos(self, doc: Document):
        """Define estilos reutilizáveis no documento."""
        # Estilo Normal base
        estilo_normal = doc.styles["Normal"]
        fonte = estilo_normal.font
        fonte.name  = "Arial"
        fonte.size  = Pt(11)
        fonte.color.rgb = TEXTO_PRIN

        # Parágrafo Normal
        pf = estilo_normal.paragraph_format
        pf.space_after  = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = Pt(16)

    # ═══════════════════════════════════════════════════════════
    # CAPA
    # ═══════════════════════════════════════════════════════════

    def _adicionar_capa(
        self, doc: Document,
        titulo: str, curso: str, uc: str, bloco: str,
        n_caps: int, professor: Optional[str]
    ):
        """Capa profissional com identidade SENAI."""
        # Quebra de seção para capa separada
        # Linha de identificação SENAI
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SENAI — SERVIÇO NACIONAL DE APRENDIZAGEM INDUSTRIAL")
        run.font.name   = "Arial"
        run.font.size   = Pt(9)
        run.font.bold   = True
        run.font.color.rgb = AZUL_MEDIO
        run.font.all_caps  = True
        p.paragraph_format.space_after = Pt(4)

        # Linha separadora
        self._linha_separadora(doc, AZUL_ESCURO)

        # Espaço
        for _ in range(6):
            doc.add_paragraph()

        # Curso
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(curso.upper())
        run.font.name      = "Arial"
        run.font.size      = Pt(12)
        run.font.bold      = True
        run.font.color.rgb = AZUL_MEDIO
        p.paragraph_format.space_after = Pt(6)

        # UC
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(uc)
        run.font.name      = "Arial"
        run.font.size      = Pt(14)
        run.font.color.rgb = TEXTO_SEC
        p.paragraph_format.space_after = Pt(24)

        # Linha decorativa
        self._linha_separadora(doc, AZUL_ESCURO, espessura=2)

        # Espaço
        for _ in range(3):
            doc.add_paragraph()

        # Título principal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.font.name      = "Arial"
        run.font.size      = Pt(24)
        run.font.bold      = True
        run.font.color.rgb = AZUL_ESCURO
        p.paragraph_format.space_after = Pt(12)

        # Bloco/Aula
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(bloco)
        run.font.name      = "Arial"
        run.font.size      = Pt(14)
        run.font.color.rgb = TEXTO_SEC
        p.paragraph_format.space_after = Pt(32)

        # Linha separadora
        self._linha_separadora(doc, AZUL_ESCURO, espessura=2)

        # Espaço
        for _ in range(6):
            doc.add_paragraph()

        # Informações de rodapé da capa
        infos = [
            f"Capítulos: {n_caps}",
            f"Versão: V1.8 — Geração Modular",
            f"Data: {datetime.now().strftime('%d/%m/%Y')}",
        ]
        if professor:
            infos.insert(0, f"Professor: {professor}")

        for info in infos:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(info)
            run.font.name      = "Arial"
            run.font.size      = Pt(10)
            run.font.color.rgb = TEXTO_SEC
            p.paragraph_format.space_after = Pt(3)

        # Quebra de página após a capa
        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SUMÁRIO
    # ═══════════════════════════════════════════════════════════

    def _adicionar_sumario(self, doc: Document, capitulos: list):
        """Sumário manual com títulos dos capítulos."""
        # Título do sumário
        p = doc.add_paragraph()
        run = p.add_run("SUMÁRIO")
        run.font.name      = "Arial"
        run.font.size      = Pt(16)
        run.font.bold      = True
        run.font.color.rgb = AZUL_ESCURO
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)

        self._linha_separadora(doc, AZUL_ESCURO)

        # Itens do sumário
        for cap in capitulos:
            num    = cap.get("numero", 0)
            titulo = cap.get("titulo", f"Capítulo {num}")
            saep   = cap.get("saep_relevante", False)

            p = doc.add_paragraph()
            run = p.add_run(f"{int(num):02d}.  {titulo}")
            run.font.name      = "Arial"
            run.font.size      = Pt(11)
            run.font.color.rgb = TEXTO_PRIN
            if saep:
                run2 = p.add_run("  [SAEP]")
                run2.font.size      = Pt(8)
                run2.font.bold      = True
                run2.font.color.rgb = AZUL_ESCURO
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.5)

            # Seções do capítulo (se disponíveis)
            for secao in cap.get("secoes", [])[:4]:
                ts = secao.get("titulo", "")
                if ts:
                    ps = doc.add_paragraph()
                    run_s = ps.add_run(f"    {ts}")
                    run_s.font.name      = "Arial"
                    run_s.font.size      = Pt(9)
                    run_s.font.color.rgb = TEXTO_SEC
                    ps.paragraph_format.space_after = Pt(1)
                    ps.paragraph_format.left_indent = Cm(1.5)

        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # CAPÍTULOS
    # ═══════════════════════════════════════════════════════════

    def _adicionar_capitulo(self, doc: Document, dados: dict):
        """Adiciona um capítulo completo ao documento."""
        numero   = dados.get("numero", 1)
        titulo   = dados.get("titulo", f"Capítulo {numero}")
        subtit   = dados.get("subtitulo", "")
        intro    = dados.get("introducao", "")
        secoes   = dados.get("secoes", [])
        resumo   = dados.get("resumo_secao", "")
        conexao  = dados.get("conexao_proximo")
        saep_rel = dados.get("saep_relevante", False)

        # ── Cabeçalho do capítulo ─────────────────────────────────────────────
        p = doc.add_paragraph()
        run = p.add_run(f"CAPÍTULO {int(numero):02d}")
        run.font.name      = "Arial"
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.color.rgb = AZUL_MEDIO
        run.font.all_caps  = True
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.font.name      = "Arial"
        run.font.size      = Pt(18)
        run.font.bold      = True
        run.font.color.rgb = AZUL_ESCURO
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)

        if subtit:
            p = doc.add_paragraph()
            run = p.add_run(subtit)
            run.font.name      = "Arial"
            run.font.size      = Pt(11)
            run.font.color.rgb = TEXTO_SEC
            run.font.italic    = True
            p.paragraph_format.space_after = Pt(8)

        self._linha_separadora(doc, AZUL_ESCURO)

        # Badge SAEP
        if saep_rel:
            p = doc.add_paragraph()
            run = p.add_run("  ALTA RELEVÂNCIA SAEP  ")
            run.font.name      = "Arial"
            run.font.size      = Pt(8)
            run.font.bold      = True
            run.font.color.rgb = AZUL_ESCURO
            p.paragraph_format.space_after = Pt(8)

        # ── Introdução do capítulo ────────────────────────────────────────────
        if intro:
            self._paragrafo(doc, intro)

        # ── Seções ───────────────────────────────────────────────────────────
        for secao in secoes:
            self._adicionar_secao(doc, secao)

        # ── Resumo do capítulo ────────────────────────────────────────────────
        if resumo:
            self._linha_separadora(doc, AZUL_CLARO)
            p = doc.add_paragraph()
            run = p.add_run("Síntese do Capítulo")
            run.font.name      = "Arial"
            run.font.size      = Pt(11)
            run.font.bold      = True
            run.font.color.rgb = AZUL_MEDIO
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.space_before = Pt(8)
            self._paragrafo(doc, resumo, italico=True)

        # ── Conexão com próximo capítulo ──────────────────────────────────────
        if conexao:
            p = doc.add_paragraph()
            run = p.add_run(f"Próximo capítulo: {conexao}")
            run.font.name      = "Arial"
            run.font.size      = Pt(9)
            run.font.italic    = True
            run.font.color.rgb = TEXTO_SEC
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.space_before = Pt(12)

        # Quebra de página entre capítulos
        doc.add_page_break()

    def _adicionar_secao(self, doc: Document, secao: dict):
        """Adiciona uma seção de conteúdo ao documento."""
        titulo_s  = secao.get("titulo", "")
        conteudo  = secao.get("conteudo", "")
        codigo    = secao.get("codigo")
        tabela    = secao.get("tabela")
        boxes     = secao.get("boxes", [])
        subsecoes = secao.get("subsecoes", [])

        # Título da seção
        if titulo_s:
            p = doc.add_paragraph()
            run = p.add_run(titulo_s)
            run.font.name      = "Arial"
            run.font.size      = Pt(13)
            run.font.bold      = True
            run.font.color.rgb = AZUL_ESCURO
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            # Borda esquerda visual simulada com espaçamento
            p.paragraph_format.left_indent = Cm(0.4)

        # Conteúdo (texto dividido em parágrafos)
        if conteudo:
            for bloco_txt in self._dividir_paragrafos(conteudo):
                if bloco_txt.strip():
                    self._paragrafo(doc, bloco_txt)

        # Código
        if codigo and isinstance(codigo, dict):
            self._adicionar_codigo(doc, codigo)

        # Tabela
        if tabela and isinstance(tabela, dict):
            self._adicionar_tabela(doc, tabela)

        # Boxes
        for box in (boxes or []):
            if isinstance(box, dict):
                self._adicionar_box(doc, box)

        # Subseções
        for sub in (subsecoes or []):
            if isinstance(sub, dict):
                self._adicionar_subsecao(doc, sub)

    def _adicionar_subsecao(self, doc: Document, sub: dict):
        """Adiciona uma subseção."""
        titulo_s = sub.get("titulo", "")
        conteudo = sub.get("conteudo", "")
        codigo   = sub.get("codigo")
        boxes    = sub.get("boxes", [])

        if titulo_s:
            p = doc.add_paragraph()
            run = p.add_run(titulo_s)
            run.font.name      = "Arial"
            run.font.size      = Pt(11)
            run.font.bold      = True
            run.font.color.rgb = AZUL_MEDIO
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.5)

        if conteudo:
            for bloco_txt in self._dividir_paragrafos(conteudo):
                if bloco_txt.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(bloco_txt.strip())
                    run.font.name      = "Arial"
                    run.font.size      = Pt(11)
                    run.font.color.rgb = TEXTO_PRIN
                    p.paragraph_format.space_after  = Pt(5)
                    p.paragraph_format.left_indent  = Cm(0.5)

        if codigo and isinstance(codigo, dict):
            self._adicionar_codigo(doc, codigo)

        for box in (boxes or []):
            if isinstance(box, dict):
                self._adicionar_box(doc, box)

    # ═══════════════════════════════════════════════════════════
    # ELEMENTOS VISUAIS
    # ═══════════════════════════════════════════════════════════

    def _adicionar_codigo(self, doc: Document, codigo: dict):
        """Bloco de código com estilo de terminal."""
        linguagem = codigo.get("linguagem", "código")
        titulo_c  = codigo.get("titulo", "")
        texto_c   = codigo.get("codigo", "")

        if not texto_c:
            return

        # Header do bloco de código
        p = doc.add_paragraph()
        label = f"{linguagem.upper()}"
        if titulo_c:
            label += f"  —  {titulo_c}"
        run = p.add_run(label)
        run.font.name      = "Courier New"
        run.font.size      = Pt(8)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.left_indent  = Cm(0.5)
        # Fundo cinza escuro via shading
        self._aplicar_shading(p, "1e293b")

        # Linhas do código
        for linha in texto_c.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(linha if linha else " ")
            run.font.name      = "Courier New"
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0xe2, 0xe8, 0xf0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.5)
            self._aplicar_shading(p, "1e293b")

        # Espaço após o código
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    def _adicionar_tabela(self, doc: Document, tabela: dict):
        """Tabela estilizada com header azul."""
        cab   = tabela.get("cabecalho", [])
        linhas = tabela.get("linhas", [])

        if not cab or not linhas:
            return

        n_cols = len(cab)
        t = doc.add_table(rows=1 + len(linhas), cols=n_cols)
        t.style = "Table Grid"

        # Cabeçalho
        hdr = t.rows[0]
        for i, texto in enumerate(cab):
            cell = hdr.cells[i]
            cell.text = texto
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(texto)
            run.font.name      = "Arial"
            run.font.size      = Pt(10)
            run.font.bold      = True
            run.font.color.rgb = BRANCO
            self._aplicar_shading_cell(cell, "0d3b66")

        # Dados
        for r_idx, linha in enumerate(linhas):
            row = t.rows[r_idx + 1]
            cor_fundo = "f5f8ff" if r_idx % 2 == 1 else "ffffff"
            for c_idx, valor in enumerate(linha[:n_cols]):
                cell = row.cells[c_idx]
                cell.text = str(valor)
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.name = "Arial"
                    p.runs[0].font.size = Pt(10)
                self._aplicar_shading_cell(cell, cor_fundo)

        # Espaço após a tabela
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _adicionar_box(self, doc: Document, box: dict):
        """Box de destaque (Atenção, Dica, SAEP, Prática)."""
        tipo     = box.get("tipo", "atencao")
        titulo_b = box.get("titulo", "")
        conteudo = box.get("conteudo", "")

        if not conteudo:
            return

        # Configuração por tipo
        config = {
            "atencao": ("d97706", "fffbeb", "ATENCAO"),
            "dica":    ("7c3aed", "f3f0ff", "DICA PROFISSIONAL"),
            "saep":    ("0d3b66", "eaf2fb", "SAEP"),
            "pratica": ("0f6b3a", "e8f5ee", "CONTEXTO PROFISSIONAL"),
        }.get(tipo, ("1a5fa8", "eaf2fb", tipo.upper()))

        cor_borda, cor_fundo, label_padrao = config
        label = titulo_b or label_padrao

        # Label do box
        p = doc.add_paragraph()
        run = p.add_run(f"  {label}  ")
        run.font.name      = "Arial"
        run.font.size      = Pt(8)
        run.font.bold      = True
        r, g, b = tuple(int(cor_borda[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.left_indent  = Cm(0.3)
        self._aplicar_shading(p, cor_fundo)

        # Conteúdo do box
        for linha in self._dividir_paragrafos(conteudo):
            if linha.strip():
                p = doc.add_paragraph()
                run = p.add_run(linha.strip())
                run.font.name      = "Arial"
                run.font.size      = Pt(10)
                run.font.color.rgb = TEXTO_PRIN
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.left_indent  = Cm(0.6)
                self._aplicar_shading(p, cor_fundo)

        # Espaço após o box
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ═══════════════════════════════════════════════════════════
    # UTILITÁRIOS
    # ═══════════════════════════════════════════════════════════

    def _paragrafo(self, doc: Document, texto: str, italico: bool = False):
        """Adiciona um parágrafo de texto normal."""
        p   = doc.add_paragraph()
        run = p.add_run(texto.strip())
        run.font.name      = "Arial"
        run.font.size      = Pt(11)
        run.font.color.rgb = TEXTO_PRIN
        run.font.italic    = italico
        p.paragraph_format.space_after   = Pt(6)
        p.paragraph_format.space_before  = Pt(0)
        p.paragraph_format.line_spacing  = Pt(16)
        return p

    def _linha_separadora(self, doc: Document, cor: RGBColor = None, espessura: int = 1):
        """Linha separadora visual."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        # RGBColor é uma tuple de 3 bytes — acessa por índice [0], [1], [2]
        if cor is not None:
            r, g, b = cor[0], cor[1], cor[2]
        else:
            r, g, b = 13, 59, 102   # AZUL_ESCURO SENAI padrão (#0d3b66)
        hex_cor = f"{r:02X}{g:02X}{b:02X}"
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(espessura * 4))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), hex_cor)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _adicionar_rodape(self, doc: Document, uc: str, bloco: str):
        """Rodapé com identidade SENAI e numeração de páginas."""
        section = doc.sections[0]
        footer  = section.footer
        p       = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(f"SENAI  |  {uc}  |  {bloco}  |  V1.8")
        run.font.name      = "Arial"
        run.font.size      = Pt(8)
        run.font.color.rgb = TEXTO_SEC

        # Número de página
        run2 = p.add_run("    Pág. ")
        run2.font.name = "Arial"
        run2.font.size = Pt(8)
        run2.font.color.rgb = TEXTO_SEC

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run3 = p.add_run()
        run3._r.append(fldChar)
        run3._r.append(instrText)
        run3._r.append(fldChar2)
        run3.font.name = "Arial"
        run3.font.size = Pt(8)
        run3.font.color.rgb = TEXTO_SEC

    def _aplicar_shading(self, paragrafo, cor_hex: str):
        """Aplica cor de fundo a um parágrafo."""
        pPr  = paragrafo._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  cor_hex.upper())
        pPr.append(shd)

    def _aplicar_shading_cell(self, cell, cor_hex: str):
        """Aplica cor de fundo a uma célula de tabela."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  cor_hex.upper())
        tcPr.append(shd)

    def _dividir_paragrafos(self, texto: str) -> list:
        """Divide texto em parágrafos por quebras de linha duplas."""
        if not texto:
            return []
        # Divide por \n\n ou \n seguido de nova linha
        partes = re.split(r'\n\n+', texto)
        resultado = []
        for parte in partes:
            # Dentro de cada parte, mantém quebras simples como espaço
            linha = parte.replace("\n", " ").strip()
            if linha:
                resultado.append(linha)
        return resultado

    def _gerar_nome(self, curso: str, uc: str, bloco: str) -> str:
        """Gera nome único para o arquivo DOCX."""
        def s(t, n):
            return re.sub(r"\s+", "_", re.sub(r"[^\w\s]", "", t).strip())[:n]
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"apostila_{s(curso,10)}_{s(uc,12)}_{s(bloco,10)}_{ts}.docx"

    def _fmt(self, b: int) -> str:
        if b < 1024:    return f"{b} B"
        if b < 1024**2: return f"{b/1024:.1f} KB"
        return f"{b/1024**2:.1f} MB"
